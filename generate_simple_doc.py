"""
Generate a simple, easy-to-read Word doc for the backend team
covering the new Registration flow and Login OTP flow for B2B Cab Portal.

Output: C:\\Users\\Pranav\\Downloads\\B2B_Registration_Login_API_Requirements.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PRIMARY = RGBColor(0x0E, 0xA5, 0xE9)   # sky-500
DARK = RGBColor(0x0F, 0x17, 0x2A)      # slate-900
GREY = RGBColor(0x64, 0x74, 0x8B)      # slate-500
LIGHT_BG = "E0F2FE"                     # sky-100
TABLE_HEADER_BG = "0EA5E9"              # sky-500
CODE_BG = "F1F5F9"                      # slate-100


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading_styled(doc, text, level=1, color=DARK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, size=11, color=DARK):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = 'Calibri'
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), CODE_BG)
    p_pr.append(shd)
    return p


def add_param_table(doc, rows):
    """rows = list of (param, type, required, description)"""
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Light Grid Accent 1'
    table.autofit = True

    headers = ['Parameter', 'Type', 'Required', 'Description']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        shade_cell(cell, TABLE_HEADER_BG)

    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            run.font.name = 'Calibri'
            if c_idx == 0:
                run.font.name = 'Consolas'
                run.bold = True

    return table


def add_endpoint_box(doc, method, url):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(13)

    method_cell = table.rows[0].cells[0]
    method_cell.text = ''
    run = method_cell.paragraphs[0].add_run(method)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'
    method_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(method_cell, "10B981" if method == "POST" else "0EA5E9")

    url_cell = table.rows[0].cells[1]
    url_cell.text = ''
    run = url_cell.paragraphs[0].add_run(url)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    shade_cell(url_cell, CODE_BG)


def add_section_break(doc):
    p = doc.add_paragraph()
    p.add_run('').font.size = Pt(4)


def main():
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ─────────── Title ───────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('B2B CAB Portal')
    run.font.size = Pt(22)
    run.font.color.rgb = PRIMARY
    run.bold = True
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Registration & Login API Requirements')
    run.font.size = Pt(14)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('For Backend Team  |  April 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    run.italic = True

    add_section_break(doc)

    # ─────────── Overview ───────────
    add_heading_styled(doc, 'Overview', level=1)
    add_para(doc,
        'We are upgrading the B2B Cab portal with a new agent registration flow '
        'and an OTP-based login option. Currently, agents can only log in with '
        'email + password. We want to add mobile number + OTP login as well, '
        'and improve the registration to be a step-by-step process with mobile '
        'verification.')

    add_para(doc, 'This document lists all the new APIs we need from the backend.', bold=True)

    add_section_break(doc)

    # ─────────── Registration Flow ───────────
    add_heading_styled(doc, 'Part 1: New Registration Flow', level=1)

    add_para(doc, 'How it works (3 steps):', bold=True)
    add_bullet(doc, 'Step 1 — Agent enters mobile number, we send OTP')
    add_bullet(doc, 'Step 2 — Agent enters OTP, we verify it and create the agent account')
    add_bullet(doc, 'Step 3 — Agent fills profile details (name, email, company, GST, etc.)')

    # 1.1 Send Registration OTP
    add_heading_styled(doc, '1.1  Send Registration OTP', level=2)
    add_para(doc, 'Send an OTP to the agent\'s mobile number for new account creation.')
    add_endpoint_box(doc, 'POST', '/auth/register-otp/send')

    add_para(doc, 'Request Parameters:', bold=True, size=10)
    add_param_table(doc, [
        ('mobile', 'string', 'Yes', '10-digit mobile number (without country code)'),
        ('countryCode', 'string', 'Yes', 'Country code, e.g. "91" for India'),
    ])

    add_para(doc, 'Response (Success):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "success",\n  "message": "OTP sent successfully",\n  "expiresIn": 300\n}')

    add_para(doc, 'Response (Error — already registered):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "error",\n  "errorCode": "MOBILE_EXISTS",\n  "message": "Mobile number already registered. Please login."\n}')

    add_para(doc, 'Notes:', bold=True, size=10)
    add_bullet(doc, 'OTP should be 6 digits and valid for 5 minutes', size=10)
    add_bullet(doc, 'If mobile is already registered as an agent, return error so we can show "Login instead"', size=10)
    add_bullet(doc, 'Apply rate limit: max 3 OTPs in 10 minutes per mobile', size=10)

    add_section_break(doc)

    # 1.2 Verify Registration OTP
    add_heading_styled(doc, '1.2  Verify Registration OTP', level=2)
    add_para(doc, 'Verify the OTP entered by the agent. If correct, create a partial agent record and return a registration token.')
    add_endpoint_box(doc, 'POST', '/auth/register-otp/verify')

    add_para(doc, 'Request Parameters:', bold=True, size=10)
    add_param_table(doc, [
        ('mobile', 'string', 'Yes', '10-digit mobile number'),
        ('countryCode', 'string', 'Yes', 'Country code, e.g. "91"'),
        ('otp', 'string', 'Yes', '6-digit OTP entered by user'),
    ])

    add_para(doc, 'Response (Success):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "success",\n  "message": "OTP verified",\n  "registrationToken": "<temp-token-valid-for-30-min>",\n  "agentId": 12345\n}')

    add_para(doc, 'Response (Error):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "error",\n  "errorCode": "INVALID_OTP",\n  "message": "Wrong OTP. Please try again."\n}')

    add_para(doc, 'Notes:', bold=True, size=10)
    add_bullet(doc, 'On success: create a partial agent record with just mobile number', size=10)
    add_bullet(doc, 'Return a temporary registration token (valid 30 minutes) — frontend will use it for the next step', size=10)
    add_bullet(doc, 'Block the account after 5 wrong attempts for 15 minutes', size=10)

    add_section_break(doc)

    # 1.3 Complete Profile
    add_heading_styled(doc, '1.3  Complete Profile', level=2)
    add_para(doc, 'Save the agent\'s profile details after OTP verification. After this, the agent is fully registered and can be logged in.')
    add_endpoint_box(doc, 'POST', '/auth/register/complete-profile')

    add_para(doc, 'Request Parameters:', bold=True, size=10)
    add_param_table(doc, [
        ('registrationToken', 'string', 'Yes', 'Token received from verify-otp step'),
        ('firstName', 'string', 'Yes', 'Agent\'s first name'),
        ('lastName', 'string', 'Yes', 'Agent\'s last name'),
        ('email', 'string', 'Yes', 'Email address (must be unique)'),
        ('password', 'string', 'Yes', 'Min 8 chars, 1 upper, 1 number'),
        ('companyName', 'string', 'Yes', 'Travel agency / company name'),
        ('city', 'string', 'Yes', 'City of operation'),
        ('gstNumber', 'string', 'No', '15-digit GSTIN (optional, can add later)'),
        ('panNumber', 'string', 'No', '10-character PAN (optional)'),
        ('referralCode', 'string', 'No', 'Referral code if any'),
    ])

    add_para(doc, 'Response (Success):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "success",\n  "message": "Registration complete",\n  "agentId": 12345,\n  "loginToken": "<JWT-for-immediate-login>",\n  "user": {\n    "id": 12345,\n    "name": "Rajesh Kumar",\n    "email": "rajesh@example.com",\n    "mobile": "9876543210"\n  }\n}')

    add_para(doc, 'Response (Error — email exists):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "error",\n  "errorCode": "EMAIL_EXISTS",\n  "message": "Email already registered. Please use another email."\n}')

    add_para(doc, 'Notes:', bold=True, size=10)
    add_bullet(doc, 'After this step the agent is fully registered and active', size=10)
    add_bullet(doc, 'Return a JWT login token so the agent can be logged in immediately (no need to login separately)', size=10)
    add_bullet(doc, 'Send a welcome email + welcome SMS to the agent', size=10)

    add_section_break(doc)

    # ─────────── Login Flow ───────────
    add_heading_styled(doc, 'Part 2: Login with Mobile OTP', level=1)

    add_para(doc, 'How it works (2 steps):', bold=True)
    add_bullet(doc, 'Step 1 — Agent enters registered mobile number, we send OTP')
    add_bullet(doc, 'Step 2 — Agent enters OTP, we verify and log them in (return JWT token)')

    add_para(doc,
        'Note: This is in addition to the existing email + password login. '
        'Both options will be available on the login page.', bold=True)

    # 2.1 Send Login OTP
    add_heading_styled(doc, '2.1  Send Login OTP', level=2)
    add_para(doc, 'Send an OTP to the agent\'s registered mobile number for login.')
    add_endpoint_box(doc, 'POST', '/auth/login-otp/send')

    add_para(doc, 'Request Parameters:', bold=True, size=10)
    add_param_table(doc, [
        ('mobile', 'string', 'Yes', 'Registered 10-digit mobile number'),
        ('countryCode', 'string', 'Yes', 'Country code, e.g. "91"'),
    ])

    add_para(doc, 'Response (Success):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "success",\n  "message": "OTP sent",\n  "expiresIn": 300\n}')

    add_para(doc, 'Response (Error — not registered):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "error",\n  "errorCode": "MOBILE_NOT_REGISTERED",\n  "message": "This mobile is not registered. Please sign up first."\n}')

    add_para(doc, 'Notes:', bold=True, size=10)
    add_bullet(doc, 'Only send OTP if the mobile belongs to an active registered agent', size=10)
    add_bullet(doc, 'Same rate limit as registration: max 3 OTPs in 10 minutes', size=10)

    add_section_break(doc)

    # 2.2 Verify Login OTP
    add_heading_styled(doc, '2.2  Verify Login OTP', level=2)
    add_para(doc, 'Verify the OTP and log the agent in. Return the same JWT token format we currently use for email/password login.')
    add_endpoint_box(doc, 'POST', '/auth/login-otp/verify')

    add_para(doc, 'Request Parameters:', bold=True, size=10)
    add_param_table(doc, [
        ('mobile', 'string', 'Yes', '10-digit mobile number'),
        ('countryCode', 'string', 'Yes', 'Country code, e.g. "91"'),
        ('otp', 'string', 'Yes', '6-digit OTP'),
    ])

    add_para(doc, 'Response (Success):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "success",\n  "message": "Login successful",\n  "loginToken": "<JWT-token-same-as-email-login>",\n  "user": {\n    "id": 12345,\n    "name": "Rajesh Kumar",\n    "email": "rajesh@example.com",\n    "mobile": "9876543210",\n    "agentId": 12345,\n    "walletBalance": 5000\n  }\n}')

    add_para(doc, 'Response (Error):', bold=True, size=10)
    add_code_block(doc, '{\n  "status": "error",\n  "errorCode": "INVALID_OTP",\n  "message": "Wrong OTP. Please try again."\n}')

    add_para(doc, 'Notes:', bold=True, size=10)
    add_bullet(doc, 'Return the EXACT same response format as the existing /user/login API so we can reuse the same frontend logic', size=10)
    add_bullet(doc, 'JWT token should have the same expiry and structure as email login token', size=10)
    add_bullet(doc, 'Block account after 5 wrong attempts for 15 minutes', size=10)

    add_section_break(doc)

    # ─────────── General Rules ───────────
    add_heading_styled(doc, 'General Rules', level=1)

    add_heading_styled(doc, 'OTP Rules', level=2)
    add_bullet(doc, 'OTP length: 6 digits, numeric only')
    add_bullet(doc, 'OTP expiry: 5 minutes from time of sending')
    add_bullet(doc, 'Rate limit: maximum 3 OTPs in 10 minutes per mobile number')
    add_bullet(doc, 'Wrong attempts: block account for 15 minutes after 5 wrong attempts')
    add_bullet(doc, 'Resend allowed only after 30 seconds (cooldown timer)')

    add_heading_styled(doc, 'SMS Provider', level=2)
    add_bullet(doc, 'Use the existing SMS gateway you use for booking confirmations')
    add_bullet(doc, 'SMS template: "Your B2B CAB OTP is {OTP}. Valid for 5 minutes. Do not share."')
    add_bullet(doc, 'Sender ID: BBCAB or similar (DLT registered)')

    add_heading_styled(doc, 'Security', level=2)
    add_bullet(doc, 'All endpoints must use HTTPS')
    add_bullet(doc, 'Hash passwords using bcrypt (cost factor 12) before storing')
    add_bullet(doc, 'Registration token (from step 1.2) should expire after 30 minutes if not used')
    add_bullet(doc, 'Login JWT should have same expiry as current email/password login')

    add_heading_styled(doc, 'Database Changes (if needed)', level=2)
    add_bullet(doc, 'Create a new table for OTPs: id, mobile, otp_hash, type (register/login), expires_at, attempts, created_at')
    add_bullet(doc, 'Add a column to agents table: registration_method ENUM("email", "mobile_otp") — to track how they signed up')
    add_bullet(doc, 'Make sure existing email/password login flow is NOT changed at all — it should keep working as is')

    add_section_break(doc)

    # ─────────── Error Codes ───────────
    add_heading_styled(doc, 'Common Error Codes', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    headers = ['Error Code', 'Meaning']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, TABLE_HEADER_BG)

    error_rows = [
        ('MOBILE_EXISTS', 'Mobile already registered (during signup)'),
        ('MOBILE_NOT_REGISTERED', 'Mobile not found in agents table (during login)'),
        ('EMAIL_EXISTS', 'Email already used by another agent'),
        ('INVALID_OTP', 'Wrong OTP entered'),
        ('OTP_EXPIRED', 'OTP is older than 5 minutes'),
        ('TOO_MANY_ATTEMPTS', 'More than 5 wrong attempts — account blocked for 15 min'),
        ('OTP_RATE_LIMIT', 'Tried to send more than 3 OTPs in 10 minutes'),
        ('INVALID_TOKEN', 'Registration token expired or invalid'),
        ('VALIDATION_ERROR', 'Required fields missing or invalid format'),
    ]

    for code, meaning in error_rows:
        row = table.add_row()
        c0 = row.cells[0]
        c0.text = ''
        run = c0.paragraphs[0].add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = DARK

        c1 = row.cells[1]
        c1.text = ''
        run = c1.paragraphs[0].add_run(meaning)
        run.font.size = Pt(9)
        run.font.color.rgb = DARK

    add_section_break(doc)

    # ─────────── Summary ───────────
    add_heading_styled(doc, 'Summary — APIs Needed', level=1)
    add_para(doc, 'Total: 5 new APIs', bold=True)
    add_bullet(doc, 'POST /auth/register-otp/send  — send OTP for signup')
    add_bullet(doc, 'POST /auth/register-otp/verify  — verify signup OTP')
    add_bullet(doc, 'POST /auth/register/complete-profile  — save profile & complete registration')
    add_bullet(doc, 'POST /auth/login-otp/send  — send OTP for login')
    add_bullet(doc, 'POST /auth/login-otp/verify  — verify login OTP & return JWT')

    add_para(doc, '')
    add_para(doc,
        'Important: Existing /user/login (email + password) API stays as is. '
        'No changes needed there.', bold=True)

    add_section_break(doc)

    # ─────────── Footer / Contact ───────────
    add_heading_styled(doc, 'Questions?', level=1)
    add_para(doc, 'For any questions about these APIs, please reach out to the frontend team.')
    add_para(doc, 'Once these APIs are ready, please share the Postman collection so we can integrate from our side.')

    # Save
    output = r"C:\Users\Pranav\Downloads\B2B_Registration_Login_API_Requirements.docx"
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == '__main__':
    main()
